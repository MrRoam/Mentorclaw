import json
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

# Define target path
target_path = Path(r"D:/Desktop/Obsidian/北航/冯如杯/冯如杯资料/范文")

result = {
    "path_checks": {
        "windows": {
            "path": r"D:\Desktop\Obsidian\北航\冯如杯\冯如杯资料\范文",
            "accessible": target_path.exists() and target_path.is_dir()
        },
        "wsl": {
            "path": "/mnt/d/Desktop/Obsidian/北航/冯如杯/冯如杯资料/范文",
            "accessible": False
        }
    },
    "docx_files": [],
    "sampled_files": []
}

if target_path.exists():
    # List docx files
    docx_files = sorted([
        f.name for f in target_path.glob("*.docx") 
        if not f.name.startswith("~$")
    ])
    result["docx_files"] = docx_files
    
    # Process first 3 files
    for filename in docx_files[:3]:
        file_path = target_path / filename
        sample = {"filename": filename, "method": None, "headings": []}
        
        try:
            # Try python-docx first
            from docx import Document
            doc = Document(str(file_path))
            headings = []
            for para in doc.paragraphs:
                style_name = para.style.name
                if 'Heading' in style_name:
                    # Extract heading level
                    level = 0
                    for char in style_name:
                        if char.isdigit():
                            level = int(char)
                            break
                    
                    text = para.text.strip()
                    if text:
                        headings.append({"level": level, "text": text})
            
            if headings or len(doc.paragraphs) > 0:
                sample["method"] = "python-docx"
                sample["headings"] = headings
        except:
            pass
        
        if not sample["method"]:
            try:
                # Fallback to XML parsing
                with zipfile.ZipFile(str(file_path), 'r') as zf:
                    xml_content = zf.read('word/document.xml')
                    root = ET.fromstring(xml_content)
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    
                    headings = []
                    for para in root.findall('.//w:p', ns):
                        # Find style
                        pStyle_elem = para.find('.//w:pStyle', ns)
                        if pStyle_elem is not None:
                            style_val = pStyle_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or pStyle_elem.get('w:val')
                            
                            if style_val and 'Heading' in style_val:
                                level = 0
                                for char in style_val:
                                    if char.isdigit():
                                        level = int(char)
                                        break
                                
                                text_elems = para.findall('.//w:t', ns)
                                text = ''.join([t.text or '' for t in text_elems]).strip()
                                
                                if text:
                                    headings.append({"level": level, "text": text})
                    
                    sample["method"] = "docx_xml_structure"
                    sample["headings"] = headings
            except:
                sample["method"] = "unavailable"
        
        result["sampled_files"].append(sample)

# Output as JSON
print(json.dumps(result, ensure_ascii=False, indent=2))
