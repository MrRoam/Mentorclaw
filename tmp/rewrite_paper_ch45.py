from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


DOC_CH4_1 = (
    "本文不再采用抽象功能项罗列的方式验证系统，而是直接以用户提供的 5 张实际运行界面截图作为证据材料。"
    "经识别，这 5 张截图共对应 3 个 demo 场景，其中两张为同一六级备考诊断界面的重复截面，因此在分析中归并为同一场景。"
    "验证重点不在于单轮回复是否“像老师”，而在于系统是否已经形成连续助学能力：能否先判断什么时候该介入，能否在介入时调到真正相关的课程资源，"
    "又能否依据学生状态给出下一步可执行安排。"
)

DOC_CH4_2 = (
    "从第一性原理看，助学智能体若想真正参与学习过程，至少要回答三个问题：什么时候帮、帮什么、如何因人而异地推进。"
    "场景一、场景二、场景三，分别对应这三类关键能力。"
)

DOC_SCENE1 = (
    "场景一：课表感知驱动的课后复盘预定。学生只表达了“基础物理没太学懂”的模糊困难，系统并没有立刻给出一段泛泛解释，"
    "而是先回到已绑定的课表，识别出基础物理学 A(1) 的固定上课时段，并明确说明会在两次上课日的当晚自动生成短复盘。"
    "更关键的是，这个复盘不是空泛提醒，而是已经细化为核心概念提炼、口语化重讲、易混点指出和 2 到 3 道小题检验。"
    "这一场景说明，梦拓龙虾已经能把“我最近听不懂”转化为“何时复盘、复盘什么、怎么检查”的明确动作。"
)

DOC_SCENE2 = (
    "场景二：课程资源驱动的复盘与作业提醒。系统识别到学生刚上完《人工智能导论》第 6 讲后，主动说明已抓取这节课的回放、PPT 和课堂笔记，"
    "并在此基础上给出两轮连续安排。第一轮是当晚的课后复盘，要求快速过核心知识点、完成 3 个理解检查题并回看没讲清楚的地方；"
    "第二轮是次日晚的短复习，要求在不看资料的情况下先回忆课堂重点，再围绕 BP、梯度下降和 CNN 等内容做针对性强化。与此同时，"
    "系统还识别出课后作业及截止时间，并把提醒前移到提交前一晚。这个场景验证的核心不是“会总结”，而是系统已经能把课程资源、复习节奏和作业节点组织成同一条执行链。"
)

DOC_SCENE3 = (
    "场景三：诊断式备考推进。原始截图中有两张内容相同的界面，均展示了系统在正式安排训练前，先询问学生当前英语水平、目标分数和主要短板。"
    "在学生给出“四级 597、目标六级 600+、写作和做题速度待提高”的回答后，系统没有套用统一模板，而是结合课内课表识别出周一晚、周三下午和周五晚三个可用时段，"
    "并为每个时段安排不同性质的训练任务：一次整套限时训练、一次错因复盘、一次写作与阅读专项强化。该场景说明，梦拓龙虾不仅能承接既有课程，"
    "也能在考试备考这类没有固定教学节奏的任务中，先做诊断，再把目标、短板和时间预算对齐。"
)

DOC_CH4_RESULT_1 = (
    "从三个场景连起来看，梦拓龙虾已经具备助学闭环的基本骨架。第一，它能够回答“什么时候介入”这个问题。"
    "场景一表明系统不再被动等待学生课后想起复习，而是能从课表直接推导出最应该承接的时间点，把复盘挂在知识刚学完、最容易遗忘之前。"
)

DOC_CH4_RESULT_2 = (
    "第二，它能够回答“此刻该帮什么”。场景二说明系统给出的支持已经不只是语言安慰，而是建立在具体课程回放、PPT、课堂笔记和作业要求上的任务化输出。"
    "这意味着梦拓龙虾的价值开始从“会说”转向“会围绕真实材料组织下一步”。"
)

DOC_CH4_RESULT_3 = (
    "第三，它能够回答“如何因人而异地推进”。场景三并没有从一开始就给出标准备考方案，而是先收集学生目标与薄弱点，再结合可用时间做安排。"
    "这说明系统的个性化并非只体现在措辞层面，而是已经开始进入节奏设计与任务拆分层面。综合来看，三个 demo 对应了时间感知、资源绑定和个性化推进三种核心能力，"
    "它们组合起来，才使系统更接近一个真正的过程型助学智能体。"
)

DOC_CH4_LIMIT_1 = (
    "首先，当前验证证据主要是流程级和界面级证据。它足以说明系统已经能组织学习过程，但还不足以单凭这几组 demo 直接证明成绩提升幅度或长期学习效果。"
)

DOC_CH4_LIMIT_2 = (
    "其次，场景三中的个性化诊断仍主要依赖学生自报信息。若要进一步提高诊断精度，后续还需要引入作业表现、阶段测验结果和更细粒度的错误记录。"
)

DOC_CH4_LIMIT_3 = (
    "再次，课表、课程资源和提醒链路虽然已经能被串起来，但多门课程并行时如何自动协调优先级、如何处理时间冲突、如何控制一周总体负荷，仍需要更强的调度策略。"
)

DOC_CH4_LIMIT_4 = (
    "最后，目前展示的主要是对话端效果。若要作为更完整的校园应用继续落地，还需要补齐通知触达、执行回执、教师或课程侧配置、以及隐私授权与数据边界等外层机制。"
)

DOC_CH5_1 = (
    "由上述三个 demo 反推，梦拓龙虾的应用前景并不在于替代教师讲授，而在于补上教学体系之外最容易断裂的那一层组织工作。"
    "对于高校课程学习，这种价值首先体现在“课后不掉线”。教师能够完成课堂讲授，但很难针对每个学生在课后继续做复盘承接、节奏提醒和作业前置督促；"
    "而梦拓龙虾恰好可以围绕课表、课程资源和个人状态，把这些最容易被忽略却最影响学习连续性的动作稳定接上。"
)

DOC_CH5_2 = (
    "第二，它适合扩展到考试备考和能力提升场景。六级备考 demo 表明，这套机制并不依赖某一门固定课程，而是可以迁移到“先诊断、再安排、再复盘”的轻量训练场景中。"
    "对于英语、计算机等级考试、保研笔试或专业资格证等任务，学生真正缺少的往往不是资料本身，而是把目标、短板和时间预算组织成持续推进链条的能力。"
    "在这一点上，梦拓龙虾具备成为个人学习教练的基础。"
)

DOC_CH5_3 = (
    "第三，它具有进一步发展为校园学习编排入口的潜力。若未来能够在授权前提下持续接入课程平台、课表、作业系统与阶段性测验结果，"
    "梦拓龙虾就不只是一个回答问题的入口，而可能成为学生管理多门课程、多类考试和阶段目标的统一支持层。其长期价值不在于把每一次回答说得更花哨，"
    "而在于让学生在繁杂任务之间始终知道：现在最该做什么，为什么先做这个，以及做完以后下一步是什么。"
)

TEX_CHAPTERS = r"""
\section{场景验证与分析}
\subsection{验证目标}
\begin{spacing}{1.5}
\setParDis
本文不再采用抽象功能项罗列的方式验证系统，而是直接以用户提供的 5 张实际运行界面截图作为证据材料。经识别，这 5 张截图共对应 3 个 demo 场景，其中两张为同一六级备考诊断界面的重复截面，因此在分析中归并为同一场景。验证重点不在于单轮回复是否“像老师”，而在于系统是否已经形成连续助学能力：能否先判断什么时候该介入，能否在介入时调到真正相关的课程资源，又能否依据学生状态给出下一步可执行安排。

从第一性原理看，助学智能体若想真正参与学习过程，至少要回答三个问题：什么时候帮、帮什么、如何因人而异地推进。场景一、场景二、场景三，分别对应这三类关键能力。
\end{spacing}

\subsection{典型场景验证}
\begin{spacing}{1.5}
\setParDis
\subsubsection{场景一：课表感知驱动的课后复盘预定}
学生只表达了“基础物理没太学懂”的模糊困难，系统并没有立刻给出一段泛泛解释，而是先回到已绑定的课表，识别出基础物理学 A(1) 的固定上课时段，并明确说明会在两次上课日的当晚自动生成短复盘。更关键的是，这个复盘不是空泛提醒，而是已经细化为核心概念提炼、口语化重讲、易混点指出和 2 到 3 道小题检验。这一场景说明，梦拓龙虾已经能把“我最近听不懂”转化为“何时复盘、复盘什么、怎么检查”的明确动作。

\begin{figure}[H]
\centering
\includegraphics[width=0.88\textwidth]{scene1_schedule_review.png}
\caption{场景一：课表感知与课后复盘安排}
\label{fig:scene1}
\end{figure}

\subsubsection{场景二：课程资源驱动的复盘与作业提醒}
系统识别到学生刚上完《人工智能导论》第 6 讲后，主动说明已抓取这节课的回放、PPT 和课堂笔记，并在此基础上给出两轮连续安排。第一轮是当晚的课后复盘，要求快速过核心知识点、完成 3 个理解检查题并回看没讲清楚的地方；第二轮是次日晚的短复习，要求在不看资料的情况下先回忆课堂重点，再围绕 BP、梯度下降和 CNN 等内容做针对性强化。与此同时，系统还识别出课后作业及截止时间，并把提醒前移到提交前一晚。这个场景验证的核心不是“会总结”，而是系统已经能把课程资源、复习节奏和作业节点组织成同一条执行链。

\begin{figure}[H]
\centering
\includegraphics[width=0.86\textwidth]{scene2_course_review_homework.png}
\caption{场景二：课程复盘与作业提醒}
\label{fig:scene2}
\end{figure}

\subsubsection{场景三：诊断式备考推进}
原始截图中有两张内容相同的界面，均展示了系统在正式安排训练前，先询问学生当前英语水平、目标分数和主要短板。在学生给出“四级 597、目标六级 600+、写作和做题速度待提高”的回答后，系统没有套用统一模板，而是结合课内课表识别出周一晚、周三下午和周五晚三个可用时段，并为每个时段安排不同性质的训练任务：一次整套限时训练、一次错因复盘、一次写作与阅读专项强化。该场景说明，梦拓龙虾不仅能承接既有课程，也能在考试备考这类没有固定教学节奏的任务中，先做诊断，再把目标、短板和时间预算对齐。

\begin{figure}[H]
\centering
\includegraphics[width=0.88\textwidth]{scene3_cet6_diagnosis_plan.png}
\caption{场景三：六级诊断与训练计划生成}
\label{fig:scene3}
\end{figure}
\end{spacing}

\subsection{结果分析}
\begin{spacing}{1.5}
\setParDis
从三个场景连起来看，梦拓龙虾已经具备助学闭环的基本骨架。第一，它能够回答“什么时候介入”这个问题。场景一表明系统不再被动等待学生课后想起复习，而是能从课表直接推导出最应该承接的时间点，把复盘挂在知识刚学完、最容易遗忘之前。

第二，它能够回答“此刻该帮什么”。场景二说明系统给出的支持已经不只是语言安慰，而是建立在具体课程回放、PPT、课堂笔记和作业要求上的任务化输出。这意味着梦拓龙虾的价值开始从“会说”转向“会围绕真实材料组织下一步”。

第三，它能够回答“如何因人而异地推进”。场景三并没有从一开始就给出标准备考方案，而是先收集学生目标与薄弱点，再结合可用时间做安排。这说明系统的个性化并非只体现在措辞层面，而是已经开始进入节奏设计与任务拆分层面。综合来看，三个 demo 对应了时间感知、资源绑定和个性化推进三种核心能力，它们组合起来，才使系统更接近一个真正的过程型助学智能体。
\end{spacing}

\subsection{局限与难点}
\begin{spacing}{1.5}
\setParDis
首先，当前验证证据主要是流程级和界面级证据。它足以说明系统已经能组织学习过程，但还不足以单凭这几组 demo 直接证明成绩提升幅度或长期学习效果。

其次，场景三中的个性化诊断仍主要依赖学生自报信息。若要进一步提高诊断精度，后续还需要引入作业表现、阶段测验结果和更细粒度的错误记录。

再次，课表、课程资源和提醒链路虽然已经能被串起来，但多门课程并行时如何自动协调优先级、如何处理时间冲突、如何控制一周总体负荷，仍需要更强的调度策略。

最后，目前展示的主要是对话端效果。若要作为更完整的校园应用继续落地，还需要补齐通知触达、执行回执、教师或课程侧配置、以及隐私授权与数据边界等外层机制。
\end{spacing}

\section{应用前景}
\begin{spacing}{1.5}
\setParDis
由上述三个 demo 反推，梦拓龙虾的应用前景并不在于替代教师讲授，而在于补上教学体系之外最容易断裂的那一层组织工作。对于高校课程学习，这种价值首先体现在“课后不掉线”。教师能够完成课堂讲授，但很难针对每个学生在课后继续做复盘承接、节奏提醒和作业前置督促；而梦拓龙虾恰好可以围绕课表、课程资源和个人状态，把这些最容易被忽略却最影响学习连续性的动作稳定接上。

第二，它适合扩展到考试备考和能力提升场景。六级备考 demo 表明，这套机制并不依赖某一门固定课程，而是可以迁移到“先诊断、再安排、再复盘”的轻量训练场景中。对于英语、计算机等级考试、保研笔试或专业资格证等任务，学生真正缺少的往往不是资料本身，而是把目标、短板和时间预算组织成持续推进链条的能力。在这一点上，梦拓龙虾具备成为个人学习教练的基础。

第三，它具有进一步发展为校园学习编排入口的潜力。若未来能够在授权前提下持续接入课程平台、课表、作业系统与阶段性测验结果，梦拓龙虾就不只是一个回答问题的入口，而可能成为学生管理多门课程、多类考试和阶段目标的统一支持层。其长期价值不在于把每一次回答说得更花哨，而在于让学生在繁杂任务之间始终知道：现在最该做什么，为什么先做这个，以及做完以后下一步是什么。
\end{spacing}
"""


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def clear_between(doc: Document, start_text: str, end_text: str) -> None:
    paragraphs = doc.paragraphs
    start_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == start_text)
    end_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == end_text)
    for idx in range(end_idx - 1, start_idx, -1):
        remove_paragraph(paragraphs[idx])


def add_text_before(marker, text: str, style: str = "Body Text") -> None:
    paragraph = marker.insert_paragraph_before(text)
    try:
        paragraph.style = style
    except KeyError:
        paragraph.style = "Normal"


def add_image_before(marker, image_path: Path, caption: str, width_cm: float = 15.2) -> None:
    image_paragraph = marker.insert_paragraph_before()
    try:
        image_paragraph.style = "Body Text"
    except KeyError:
        image_paragraph.style = "Normal"
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))

    caption_paragraph = marker.insert_paragraph_before(caption)
    try:
        caption_paragraph.style = "Body Text"
    except KeyError:
        caption_paragraph.style = "Normal"
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def rewrite_docx(src: Path, dst: Path, paper_dir: Path) -> None:
    shutil.copy2(src, dst)
    doc = Document(dst)

    clear_between(doc, "（一）验证目标", "（二）典型场景验证")
    clear_between(doc, "（二）典型场景验证", "（三）结果分析")
    clear_between(doc, "（三）结果分析", "（四）局限与难点")
    clear_between(doc, "（四）局限与难点", "五、应用前景")
    clear_between(doc, "五、应用前景", "结论")

    marker = find_paragraph(doc, "（二）典型场景验证")
    add_text_before(marker, DOC_CH4_1)
    add_text_before(marker, DOC_CH4_2)

    marker = find_paragraph(doc, "（三）结果分析")
    add_text_before(marker, DOC_SCENE1)
    add_image_before(
        marker,
        paper_dir / "配图" / "scene1_schedule_review.png",
        "图 3  场景一：课表感知与课后复盘安排",
    )
    add_text_before(marker, DOC_SCENE2)
    add_image_before(
        marker,
        paper_dir / "配图" / "scene2_course_review_homework.png",
        "图 4  场景二：课程复盘与作业提醒",
    )
    add_text_before(marker, DOC_SCENE3)
    add_image_before(
        marker,
        paper_dir / "配图" / "scene3_cet6_diagnosis_plan.png",
        "图 5  场景三：六级诊断与训练计划生成",
        width_cm=15.0,
    )

    marker = find_paragraph(doc, "（四）局限与难点")
    add_text_before(marker, DOC_CH4_RESULT_1)
    add_text_before(marker, DOC_CH4_RESULT_2)
    add_text_before(marker, DOC_CH4_RESULT_3)

    marker = find_paragraph(doc, "五、应用前景")
    add_text_before(marker, DOC_CH4_LIMIT_1)
    add_text_before(marker, DOC_CH4_LIMIT_2)
    add_text_before(marker, DOC_CH4_LIMIT_3)
    add_text_before(marker, DOC_CH4_LIMIT_4)

    marker = find_paragraph(doc, "结论")
    add_text_before(marker, DOC_CH5_1)
    add_text_before(marker, DOC_CH5_2)
    add_text_before(marker, DOC_CH5_3)

    doc.save(dst)


def ensure_template_assets(docx_src: Path, paper_dir: Path) -> None:
    include_dir = paper_dir / "include_picture"
    include_dir.mkdir(exist_ok=True)

    with zipfile.ZipFile(docx_src) as zf:
        media_map = {
            "word/media/image1.png": include_dir / "xiaohui.png",
            "word/media/image2.png": include_dir / "xiaoming.png",
            "word/media/image3.jpeg": include_dir / "image3.jpeg",
        }
        for src_name, dst_path in media_map.items():
            if src_name in zf.namelist():
                dst_path.write_bytes(zf.read(src_name))

    architecture_candidates = sorted((paper_dir / "配图").glob("*总体功能架构图*.png"))
    if architecture_candidates:
        shutil.copy2(architecture_candidates[0], include_dir / "mengtuo_architecture.png")

    for name in [
        "scene1_schedule_review.png",
        "scene2_course_review_homework.png",
        "scene3_cet6_diagnosis_plan.png",
    ]:
        src = paper_dir / "配图" / name
        if src.exists():
            shutil.copy2(src, include_dir / name)


def rewrite_tex(src: Path, dst: Path) -> None:
    text = src.read_text(encoding="utf-8")
    start = text.index(r"\section{场景验证与分析}")
    end = text.index(r"\section*{结论}")
    new_text = text[:start] + TEX_CHAPTERS + "\n" + text[end:]
    dst.write_text(new_text, encoding="utf-8")


def main() -> None:
    paper_dir = Path(r"D:\Desktop\Obsidian\北航\冯如杯\2稿\梦拓龙虾论文成稿")
    docx_src = paper_dir / "梦拓龙虾-论文-大学生助学版-修复.docx"
    docx_dst = paper_dir / "梦拓龙虾-论文-大学生助学版-第四第五章重写.docx"
    tex_src = paper_dir / "梦拓龙虾-论文-v2.tex"
    tex_dst = paper_dir / "梦拓龙虾-论文-v2-第四第五章重写.tex"

    ensure_template_assets(docx_src, paper_dir)
    rewrite_docx(docx_src, docx_dst, paper_dir)
    rewrite_tex(tex_src, tex_dst)

    print(docx_dst)
    print(tex_dst)


if __name__ == "__main__":
    main()
