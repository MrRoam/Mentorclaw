import sys
import json
from pathlib import Path

# Try to import docx
try:
    from docx import Document
    print("python-docx is available", file=sys.stderr)
except ImportError:
    print("python-docx not available", file=sys.stderr)
    sys.exit(1)

# Define path
target_path = Path("D:/Desktop/Obsidian/北航/冯如杯/冯如杯资料/范文")

if not target_path.exists():
    print(f"Path not found: {target_path}", file=sys.stderr)
    sys.exit(1)

print(f"Path exists: {target_path}", file=sys.stderr)

# List docx files
docx_files = sorted([f.name for f in target_path.glob("*.docx") if not f.name.startswith("~$")])
print(f"Found {len(docx_files)} files: {docx_files}", file=sys.stderr)

# Process first 3
for filename in docx_files[:3]:
    file_path = target_path / filename
    print(f"\nProcessing: {filename}", file=sys.stderr)
    
    try:
        doc = Document(str(file_path))
        print(f"  - Document opened successfully", file=sys.stderr)
        print(f"  - Total paragraphs: {len(doc.paragraphs)}", file=sys.stderr)
        
        headings = []
        for i, para in enumerate(doc.paragraphs[:20]):  # Check first 20
            style_name = para.style.name
            print(f"    Para {i}: style='{style_name}', text='{para.text[:50] if para.text else 'EMPTY'}'", file=sys.stderr)
            
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.split()[-1])
                except:
                    level = 0
                if para.text.strip():
                    headings.append({"level": level, "text": para.text.strip()})
        
        print(f"  - Found {len(headings)} headings", file=sys.stderr)
        
    except Exception as e:
        print(f"  - Error: {e}", file=sys.stderr)
