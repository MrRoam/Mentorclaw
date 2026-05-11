from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, cn="宋体", en="Times New Roman", size=Pt(12), bold=False):
    run.font.name = en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = size
    run.font.bold = bold


def add_text(paragraph, text, cn="宋体", en="Times New Roman", size=Pt(12), bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, cn=cn, en=en, size=size, bold=bold)
    return run


def add_paragraph(
    doc,
    text="",
    *,
    cn="宋体",
    en="Times New Roman",
    size=Pt(12),
    bold=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=True,
    line_spacing=1.5,
    space_before=0,
    space_after=0,
):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if first_line_indent:
        fmt.first_line_indent = Cm(0.74)
    add_text(p, text, cn=cn, en=en, size=size, bold=bold)
    return p


def add_heading_paragraph(doc, text, level):
    if level == 1:
        return add_paragraph(
            doc,
            text,
            cn="黑体",
            size=Pt(16),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            line_spacing=1.0,
            space_before=12,
            space_after=6,
        )
    if level == 2:
        return add_paragraph(
            doc,
            text,
            cn="黑体",
            size=Pt(14),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=False,
            line_spacing=1.0,
            space_before=8,
            space_after=4,
        )
    return add_paragraph(
        doc,
        text,
        cn="黑体",
        size=Pt(12),
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        line_spacing=1.0,
        space_before=6,
        space_after=2,
    )


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text_run = paragraph.add_run("1")
    set_run_font(text_run, cn="Times New Roman", en="Times New Roman", size=Pt(10.5), bold=False)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    text_run._r.append(fld_end)


def set_page_number_format(section, start=1, fmt="decimal"):
    sectPr = section._sectPr
    elems = sectPr.xpath("./w:pgNumType")
    elem = elems[0] if elems else OxmlElement("w:pgNumType")
    elem.set(qn("w:start"), str(start))
    elem.set(qn("w:fmt"), fmt)
    if not elems:
        sectPr.append(elem)


def set_margins(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)


doc = Document()

for sec in doc.sections:
    set_margins(sec)

# Cover
for _ in range(4):
    add_paragraph(doc, "", first_line_indent=False, line_spacing=1.0)

add_paragraph(
    doc,
    '第三十六届“冯如杯”创意赛道项目论文',
    cn="华文中宋",
    size=Pt(22),
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=False,
    line_spacing=1.0,
)
for _ in range(3):
    add_paragraph(doc, "", first_line_indent=False, line_spacing=1.0)

add_paragraph(
    doc,
    "梦拓龙虾",
    cn="华文中宋",
    size=Pt(22),
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=False,
    line_spacing=1.0,
)

add_paragraph(
    doc,
    "——面向高校课程学习闭环的助学智能体设计与实现",
    cn="华文新魏",
    size=Pt(16),
    bold=False,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    first_line_indent=False,
    line_spacing=1.0,
)

for _ in range(14):
    add_paragraph(doc, "", first_line_indent=False, line_spacing=1.0)

# Section 2: abstract, no header, roman page numbers
section2 = doc.add_section(WD_SECTION.NEW_PAGE)
set_margins(section2)
section2.header.is_linked_to_previous = False
section2.footer.is_linked_to_previous = False
section2.header.paragraphs[0].text = ""
section2.footer.paragraphs[0].clear()
set_page_number_format(section2, start=1, fmt="lowerRoman")
add_page_number(section2.footer.paragraphs[0])

add_heading_paragraph(doc, "摘要", 1)
add_paragraph(
    doc,
    "高校学生在课程学习中普遍面临资料分散、任务组织零散、课后复盘容易中断、复习节奏难以长期维持等问题。现有学习平台更偏重资源存放与流程管理，通用聊天式 AI 虽然能够进行问答与总结，但难以持续理解学生在具体课程中的学习状态，更难主动推动学习过程。针对这一问题，本文提出助学智能体“梦拓龙虾”。该作品面向高校课程学习场景，以课程为核心组织单元，围绕“目标建立—任务推进—课后复盘—周期复习—动态调整”构建学习闭环支持机制。系统在 OpenClaw 运行环境中完成原型搭建，并由项目自主实现课程状态维护、学习任务编排、课程资料聚合调用、课后总结、复习提醒与个性化主动任务等核心功能。通过典型场景验证，梦拓龙虾能够在课程学习中较好承担学习组织者和推进者的角色，帮助学生减少在资料查找、任务衔接和复习坚持方面的额外负担。本文进一步分析了该方案的创意价值、技术可行性与应用前景，并指出其在学习效果量化、自动化执行深度和多课程协同等方面仍有提升空间。",
)
add_paragraph(
    doc,
    "关键词：助学智能体，课程学习闭环，学习规划，课后复盘，主动学习支持",
    cn="宋体",
    size=Pt(12),
    bold=True,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent=False,
    line_spacing=1.5,
)

add_heading_paragraph(doc, "Abstract", 1)
add_paragraph(
    doc,
    "College students often face fragmented learning materials, discontinuous post-class review, and difficulty maintaining a stable review rhythm across courses. Conventional learning platforms focus on storing resources, while general conversational AI mainly supports single-turn question answering and lacks sustained understanding of a learner's course-specific progress. To address this gap, this paper presents Mentor Lobster, an educational agent designed for the closed loop of university course learning. The system takes courses as the core organizational unit and supports goal setting, task progression, post-class reflection, periodic review, and dynamic adjustment. Built as a prototype in the OpenClaw runtime, the project independently implements course-state maintenance, learning-task orchestration, resource aggregation, post-class summarization, review prompting, and personalized proactive task support. Scenario-based validation shows that the system can effectively help students organize learning processes, reduce friction in resource retrieval and task continuity, and improve the sustainability of course review. The paper also discusses the creativity, feasibility, and future application potential of the project, while noting current limitations in evaluation depth, automation maturity, and multi-course coordination.",
    cn="Times New Roman",
    en="Times New Roman",
    size=Pt(12),
    bold=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=False,
    line_spacing=1.5,
)
add_paragraph(
    doc,
    "Keywords: educational agent, course learning loop, learning planning, post-class reflection, proactive learning support",
    cn="Times New Roman",
    en="Times New Roman",
    size=Pt(12),
    bold=True,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent=False,
    line_spacing=1.5,
)

# Section 3: body with header and arabic page numbers
section3 = doc.add_section(WD_SECTION.NEW_PAGE)
set_margins(section3)
section3.header.is_linked_to_previous = False
section3.footer.is_linked_to_previous = False
set_page_number_format(section3, start=1, fmt="decimal")

hp = section3.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(hp, '北京航空航天大学第三十六届“冯如杯”竞赛创意赛道参赛作品', cn="宋体", en="Times New Roman", size=Pt(9), bold=False)
section3.footer.paragraphs[0].clear()
add_page_number(section3.footer.paragraphs[0])

add_heading_paragraph(doc, "一、引言", 1)
add_heading_paragraph(doc, "（一）研究背景与创意来源", 2)
for para in [
    "在高校课程学习中，学生真正遇到的困难往往不是“不会某一道题”，而是“如何把一门课持续学下去”。一门课程通常包含课表、通知、作业、课件、录播、课堂重点和考试要求等多类信息，这些信息分散在多个平台和时间节点中。学生即使在某一时刻知道自己该做什么，也常常因为任务切换频繁、信息入口分散和缺少持续提醒而中断学习节奏。",
    "本项目的创意来源正是对这一现实问题的观察。我们认为，课程学习不是一个个孤立问题的集合，而是一个需要持续组织的过程。如果一个系统只能在学生提问时回答问题，却不能帮助学生建立目标、衔接任务、复盘课程和推进复习，那么它对学习效率的提升仍然有限。",
    "因此，梦拓龙虾并不把自己定位为普通问答助手，而是定位为面向课程学习过程的助学智能体。它希望扮演的角色不是“替学生学习”，而是“帮助学生把学习过程组织起来”。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（二）现有方案分析", 2)
for para in [
    "现有学习平台的优势在于规范化和结构化，能够较稳定地承载课表、通知、作业和资源，但其主要作用仍是信息发布和流程管理。学生可以在平台中看到大量内容，却仍然需要自己完成筛选、整理、衔接和推进。",
    "通用聊天式 AI 的优势在于表达能力和即时交互。它能够解释概念、总结资料、改写文本，甚至帮助制定计划。但在课程学习场景中，它往往停留在“问一次、答一次”的交互模式，难以持续维护学生在某门课中的目标、进度和薄弱点。",
    "从学习科学的角度看，学习效果不仅取决于知识解释是否清楚，还取决于是否存在持续的组织、反馈和复习机制。个性化支持长期被认为能够显著提高学习质量[1]，而分散复习和提取练习对长期保持也具有明确作用[2][3]。因此，真正适合课程学习的系统，应同时具备状态持续、任务推进和复习支持三种能力。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（三）本文工作", 2)
for para in [
    "针对上述问题，本文提出并设计了助学智能体“梦拓龙虾”。与强调单次问答效果的方案不同，本作品围绕高校课程学习闭环展开，重点解决课程资料分散、学习过程缺少外部组织、课后复盘和复习难以持续的问题。",
    "本文的主要工作包括：一是提出以课程为核心组织单元的助学智能体设计思路；二是围绕目标建立、任务推进、课后总结、复习提醒和个性化主动任务设计完整的功能闭环；三是基于原型系统对典型学习场景进行验证，并分析其应用价值和后续改进方向。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "二、作品核心创意与总体方案", 1)
add_heading_paragraph(doc, "（一）设计目标与系统定位", 2)
for para in [
    "梦拓龙虾的设计目标不是成为一个“大而全”的生活管理工具，而是成为一个围绕课程学习持续运行的助学智能体。其核心目标可以概括为三点：帮助学生快速看清当前课程目标，帮助学生持续维持学习节奏，帮助学生在课后形成可继续推进的复习链条。",
    "围绕这一目标，梦拓龙虾将服务边界聚焦在课程相关场景，包括课程目标建立、阶段任务拆解、课堂资料调用、课后总结、复习提醒以及少量与课程推进直接相关的主动任务。这样的定位使作品能够把有限能力集中到最影响学习体验的环节。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（二）课程学习闭环模型", 2)
for para in [
    "梦拓龙虾的核心创意可以概括为“把一门课从一次次零散提问，转化为一个持续推进的学习闭环”。该闭环包括五个核心环节：目标建立、任务推进、课后复盘、周期复习和动态调整。",
    "其中，目标建立回答“这门课当前要达成什么”；任务推进回答“此刻应该先做什么”；课后复盘回答“刚学完的内容如何沉淀”；周期复习回答“哪些知识点需要再次回看”；动态调整回答“当时间、进度或理解情况变化时如何及时修正路径”。",
]:
    add_paragraph(doc, para)
add_paragraph(
    doc,
    "图1  梦拓龙虾课程学习闭环流程图（图片待补）",
    cn="宋体",
    size=Pt(10.5),
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=False,
    line_spacing=1.0,
    space_before=6,
    space_after=6,
)

add_heading_paragraph(doc, "（三）总体功能架构", 2)
for para in [
    "从功能上看，梦拓龙虾主要由四部分构成：学习状态维护模块、学习任务编排模块、课程资料聚合模块和主动支持模块。学习状态维护模块负责记录当前课程目标、阶段进度和学习偏好；学习任务编排模块负责把模糊目标拆解成可执行的下一步；课程资料聚合模块负责围绕当前问题调用课表、课件、录播和作业等信息；主动支持模块负责在合适的时机触发课后总结、复习提醒和阶段检查。",
    "四个模块并不是彼此孤立，而是共同围绕课程学习过程服务。状态维护提供连续性，任务编排提供方向感，资料聚合提供依据，主动支持则保证学习过程不被轻易打断。",
]:
    add_paragraph(doc, para)
add_paragraph(
    doc,
    "图2  梦拓龙虾总体功能架构图（图片待补）",
    cn="宋体",
    size=Pt(10.5),
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=False,
    line_spacing=1.0,
    space_before=6,
    space_after=6,
)

add_heading_paragraph(doc, "（四）关键运行逻辑", 2)
for para in [
    "当学生输入新的学习目标或课程问题后，系统首先识别当前需求属于规划、辅导、复盘还是提醒；随后调用与当前课程有关的最小必要资料，形成围绕当前任务的支持内容；在本轮交互结束后，再把关键信息更新回课程状态中，以便下一次继续衔接。",
    "这一运行逻辑的重点不在于“尽量多说”，而在于“尽量让下一步明确”。从评委视角看，这也是作品的核心价值所在：它解决的不是单一回答能力，而是学习过程连续性问题。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "三、关键功能设计与可行性分析", 1)
add_heading_paragraph(doc, "（一）长期学习状态维护机制", 2)
for para in [
    "课程学习之所以容易中断，很大程度上是因为目标和进度没有稳定保留下来。梦拓龙虾通过长期学习状态维护机制，记录课程范围、当前目标、阶段任务、学习偏好和需要重点关注的问题，使系统能够在下一次交互中直接续接，而不必每次从头开始。",
    "这种机制并不追求建立复杂而抽象的学习画像，而是优先保存真正对课程推进有帮助的信息，例如正在准备哪门课、下一步要完成什么、哪些内容需要复盘。这样既能减少冗余，也更符合实际使用场景。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（二）学习计划生成与动态调整", 2)
for para in [
    "学生在输入目标时，往往给出的只是模糊表述，如“准备期中考试”“这周补完这门课”。梦拓龙虾会先将这些目标转化为更清晰的课程任务，再根据当前阶段形成可执行的下一步，而不是直接输出一份看起来完整却难以落实的大计划。",
    "当学习进度、时间安排或理解情况发生变化时，系统能够重新调整任务顺序和关注重点。这种动态调整能力使作品更贴近真实学习过程，因为课程学习从来不是一条完全静态的路径。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（三）课程资料聚合与定向调用", 2)
for para in [
    "课程学习支持的关键在于“知道该从哪里找材料”。梦拓龙虾围绕课程场景整合课表、通知、课件、录播、作业和相关资料，使学生不必在多个平台之间频繁切换。系统并不是简单堆叠资料，而是根据当前任务调用最相关的内容，从而降低信息负担。",
    "这一设计使作品兼具实用性与可扩展性。一方面，它能直接服务于课堂学习和课后复盘；另一方面，随着课程资源接入范围扩大，系统可以继续增强对不同课程和不同资料类型的支持能力。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（四）课后总结与周期复习支持", 2)
for para in [
    "很多学生在课堂结束后并不是没有意识到复习重要，而是难以把“刚学过的内容”及时转化为“之后还能继续回看的重点”。梦拓龙虾在每次课程结束后可提供课后总结，并在后续适当时间提醒学生回顾核心知识点、待理解概念和后续任务。",
    "这种支持方式与分散复习和提取练习的学习规律是一致的[2][3]。相比把复习完全交给学生自觉管理，系统能够在合适时机提供更低成本的推动，帮助学生维持学习连续性。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（五）个性化主动任务机制", 2)
for para in [
    "梦拓龙虾不仅响应学生主动提问，也允许学生预设与课程相关的主动任务，例如课后自动总结、每周复习提醒、阶段检查等。其意义在于把学习支持从“等学生想起来再做”变成“在需要的时候主动出现”。",
    "对创意赛道作品而言，这一设计体现了作品的主动性与成长性。它意味着作品未来不仅可以做答疑，还可以成为一个真正参与学习组织过程的智能体。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（六）实现基础与技术可行性", 2)
for para in [
    "从技术实现角度看，梦拓龙虾并非停留在概念层面，而是已经形成可运行原型。项目以 OpenClaw 作为运行环境，并自主实现了面向课程学习的状态管理、任务编排和资料调用模块。课程资料的组织、任务推进逻辑以及主动任务机制均具有明确的实现路径。",
    "从工程可行性看，当前方案的实现依赖的并不是罕见硬件或不可获取资源，而是现有的大模型能力、课程平台数据接口和常规的状态管理方法。因此，该作品具备进一步落地和持续完善的现实基础。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "四、典型应用场景与效果分析", 1)
add_heading_paragraph(doc, "（一）场景一：新课程目标创建", 2)
for para in [
    "当学生输入“我要在两周内完成某门课的期中复习”时，梦拓龙虾能够围绕这门课建立目标，帮助学生明确复习范围、当前优先级和第一步行动。相比普通计划清单，它更强调课程上下文和后续可持续推进。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（二）场景二：课后复盘与资料调用", 2)
for para in [
    "在课后复盘场景中，学生通常最需要的是快速回到真正有用的资料，而不是重新翻遍所有平台。梦拓龙虾能够围绕当前课程任务调用相关课件、录播和作业信息，并以更清晰的方式整理出来，减少学生在信息查找上的额外消耗。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（三）场景三：阶段复习与提醒", 2)
for para in [
    "在阶段复习场景中，系统能够根据课程进度和既有目标，在适当时间提醒学生回顾重点、检查完成情况并调整下一步计划。这种持续提醒机制尤其适合任务多、节奏快的高校课程环境。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（四）综合价值分析", 2)
for para in [
    "从评委视角看，梦拓龙虾的价值不在于“又做了一个聊天机器人”，而在于它抓住了高校课程学习中真正高频、却长期缺少好工具支持的问题：学习过程组织。它把学习支持从单次问答扩展到目标建立、过程推进和复习维持，因而更接近真实教育场景。",
    "从用户价值看，作品可以减少资料切换成本、降低拖延和中断概率，并帮助学生把零散学习活动组织成连续过程。这种价值并不依赖极端理想化的使用条件，因此具备较好的普适性。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "（五）当前不足与改进方向", 2)
for para in [
    "目前，梦拓龙虾仍处于原型阶段。一方面，学习效果评估还偏初步，对学生理解深度和长期掌握程度的量化仍需加强；另一方面，主动任务机制虽然已经具备明确思路，但在执行丰富度和多课程协同上仍有进一步提升空间。",
    "后续改进可以围绕三条主线展开：一是增强对学习成效的判断能力；二是完善主动任务的执行链路；三是进一步优化多门课程并行时的优先级管理能力。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "五、应用前景", 1)
for para in [
    "梦拓龙虾的应用前景主要体现在高校日常课程学习、阶段性考试复习和自主学习督导三个方面。对于高校课程而言，它可以作为面向日常学习的助理；对于考试复习而言，它可以帮助学生维持更连续的复习节奏；对于自主学习而言，它可以充当稳定的过程组织者。",
    "随着课程资料接入能力和主动任务机制的完善，梦拓龙虾还可以进一步扩展到跨课程协同、长期学习陪伴和更细粒度的个性化支持。由此看，该作品具有较好的推广空间和持续演化潜力。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "结论", 1)
for para in [
    "本文从高校课程学习的真实问题出发，提出并设计了助学智能体“梦拓龙虾”。该作品的核心创意不在于强化单轮问答能力，而在于围绕课程学习过程构建一个能够持续组织、持续推进和持续提醒的智能支持机制。",
    "通过对目标建立、任务推进、课后复盘、周期复习和主动任务等环节的设计与分析，可以看出梦拓龙虾在高校课程场景中具有明确的应用价值和较强的现实意义。未来，随着功能进一步完善，该作品有望成为面向高校学习场景的长期陪伴式智能工具。",
]:
    add_paragraph(doc, para)

add_heading_paragraph(doc, "参考文献", 1)
refs = [
    "[1] Bloom B S. The 2 Sigma Problem: The Search for Methods of Group Instruction as Effective as One-to-One Tutoring[J]. Educational Researcher, 1984, 13(6): 4-16.",
    "[2] Cepeda N J, Pashler H, Vul E, et al. Distributed Practice in Verbal Recall Tasks: A Review and Quantitative Synthesis[J]. Psychological Bulletin, 2006, 132(3): 354-380.",
    "[3] Karpicke J D, Blunt J R. Retrieval Practice Produces More Learning than Elaborative Studying with Concept Mapping[J]. Science, 2011, 331(6018): 772-775.",
    "[4] Dunlosky J, Rawson K A, Marsh E J, et al. Improving Students' Learning With Effective Learning Techniques[J]. Psychological Science in the Public Interest, 2013, 14(1): 4-58.",
    "[5] Corbett A T, Anderson J R. Knowledge Tracing: Modeling the Acquisition of Procedural Knowledge[J]. User Modeling and User-Adapted Interaction, 1994, 4(4): 253-278.",
    "[6] Park J S, O'Brien J C, Cai C J, et al. Generative Agents: Interactive Simulacra of Human Behavior[EB/OL]. arXiv:2304.03442, 2023.",
]
for ref in refs:
    add_paragraph(
        doc,
        ref,
        cn="宋体",
        en="Times New Roman",
        size=Pt(10.5),
        bold=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        line_spacing=1.5,
    )

out = r"/home/jiaxu/mentorclaw-source/output/doc/梦拓龙虾-论文-Word版.docx"
doc.save(out)
print(out)
